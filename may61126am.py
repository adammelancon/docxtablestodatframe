from docx import Document
import pandas as pd
import os
import sys
import feather
import openpyxl

#document = Document('./librefiles/2004.docx')

biglist =[]
path = './librefiles/'
worddocs_list = []
#print(os.listdir(path))

for filename in os.listdir(path):
    wordDoc = Document(path + filename)
    worddocs_list.append(wordDoc)

#print(worddocs_list)
filenumber = 0

for wordDoc in worddocs_list:
    filenumber += 1
    print("\n" + "Getting file" + " " + str(filenumber))
    for table in wordDoc.tables:
        doctbls=[]
        tbllist=[]
        rowlist=[]
        sys.stdout.write(str(". "))
        sys.stdout.flush()
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                #print(rowlist)
                rowlist.append(cell.text)
            tbllist.append(rowlist)
            rowlist=[]
        doctbls = doctbls + tbllist
        biglist = biglist + doctbls

print("Making the Dataframe!")
df = pd.DataFrame(biglist)     
print("Moving Header to the Top!")
df.rename(columns=df.iloc[0], inplace = True)
print("Removing Duplicate Headers")
df = df[df['NAME'] != 'NAME']
print("Printing Dataframe")
print(df)
print("Writing Dataframe to Feather File")
feather.write_dataframe(df, './FINALOBITMERGE.feather')
print("Writing Dataframe to Excel File")
df.to_excel("./FINALOBITMERGE.xlsx", sheet_name="Obits2004-2011")
print("END")